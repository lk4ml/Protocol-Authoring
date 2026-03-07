"""
ICH M11 protocol document generator.

Produces a Word (.docx) document from a ``ProtocolDB`` database row using
the ``python-docx`` library.  The generated document follows the ICH M11
recommended structure with the following sections:

    - Title Page
    - Protocol Synopsis (summary table)
    - Schedule of Activities (encounters x activities matrix)
    - Section 1: Introduction
    - Section 2: Objectives and Endpoints
    - Section 3: Trial Design
    - Section 4: Trial Population
    - Section 5: Trial Interventions
    - Section 6: Endpoints
    - Sections 7-11: Additional narrative sections
"""

from __future__ import annotations

from datetime import datetime, timezone
from typing import Any, TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if TYPE_CHECKING:
    from docx.table import Table as DocxTable


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell: Any, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table: "DocxTable") -> None:
    """Apply thin black borders to every cell edge in *table*."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def _add_styled_paragraph(document: Document, text: str, style: str = "Normal",
                          bold: bool = False, alignment: int | None = None,
                          font_size: int | None = None,
                          space_after: int | None = None) -> Any:
    """Add a paragraph with optional inline formatting overrides."""
    para = document.add_paragraph(style=style)
    run = para.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if alignment is not None:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _add_title_page(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Render the protocol title page."""
    doc.add_paragraph()  # spacer

    _add_styled_paragraph(doc, "CLINICAL TRIAL PROTOCOL",
                          bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          font_size=18, space_after=6)

    _add_styled_paragraph(doc, protocol_db.protocol_number or "[Protocol Number]",
                          bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          font_size=14, space_after=12)

    _add_styled_paragraph(doc, protocol_db.full_title or "[Full Title]",
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          font_size=12, space_after=24)

    # Key metadata table
    rows_data = [
        ("Sponsor:", protocol_db.sponsor_name or ""),
        ("Phase:", protocol_db.phase or ""),
        ("Study Type:", protocol_db.study_type or ""),
        ("Protocol Version:", protocol_db.version or "1.0"),
        ("Date:", datetime.now(timezone.utc).strftime("%d %B %Y")),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value
        for para in cell_label.paragraphs:
            for run in para.runs:
                run.bold = True

    doc.add_page_break()


def _add_synopsis(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Render the Protocol Synopsis section as a summary table."""
    doc.add_heading("PROTOCOL SYNOPSIS", level=1)

    study_design_data: dict = protocol_db.study_design_data or {}

    # Collect arms description
    arms = study_design_data.get("studyArms", [])
    arms_text = ", ".join(a.get("name", "") for a in arms) if arms else "Not specified"

    # Collect objectives summary
    objectives = study_design_data.get("objectives", [])
    def _level_decode(obj):
        lvl = obj.get("level", "")
        if isinstance(lvl, dict):
            return lvl.get("decode", "")
        return str(lvl)

    primary_objs = [o.get("name", "") for o in objectives
                    if _level_decode(o).lower() == "primary"]
    primary_text = "; ".join(primary_objs) if primary_objs else "Not specified"

    # Sample size
    sample_size = protocol_db.sample_size
    sample_text = str(sample_size) if sample_size else "Not specified"

    # Intervention model
    im = study_design_data.get("interventionModel")
    im_text = im.get("decode", "Not specified") if isinstance(im, dict) else "Not specified"

    # Blinding
    bs = study_design_data.get("blindingSchema")
    bs_text = bs.get("decode", "Not specified") if isinstance(bs, dict) else "Not specified"

    rows_data = [
        ("Protocol Number", protocol_db.protocol_number or ""),
        ("Full Title", protocol_db.full_title or ""),
        ("Short Title", protocol_db.short_title or ""),
        ("Sponsor", protocol_db.sponsor_name or ""),
        ("Phase", protocol_db.phase or ""),
        ("Study Type", protocol_db.study_type or ""),
        ("Study Design", im_text),
        ("Blinding", bs_text),
        ("Number of Arms", str(len(arms)) if arms else "0"),
        ("Arms", arms_text),
        ("Primary Objective(s)", primary_text),
        ("Planned Sample Size", sample_text),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    _set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (field, value) in enumerate(rows_data):
        cell_field = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_field.text = field
        cell_value.text = value
        _set_cell_shading(cell_field, "D9E2F3")
        for para in cell_field.paragraphs:
            for run in para.runs:
                run.bold = True

    doc.add_paragraph()  # spacer
    doc.add_page_break()


def _add_schedule_of_activities(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Render the Schedule of Activities (SoA) matrix.

    Encounters form the columns; activities form the rows.  A cell
    contains ``X`` for mandatory instances and ``C`` for conditional.
    """
    doc.add_heading("SCHEDULE OF ACTIVITIES", level=1)

    study_design_data: dict = protocol_db.study_design_data or {}
    encounters: list[dict] = study_design_data.get("encounters", [])
    activities: list[dict] = study_design_data.get("activities", [])
    timelines: list[dict] = study_design_data.get("scheduleTimelines", [])

    if not encounters or not activities:
        doc.add_paragraph("No schedule of activities data available.")
        doc.add_page_break()
        return

    # Sort by order
    encounters = sorted(encounters, key=lambda e: e.get("order", 0))
    encounter_ids = [e.get("id") for e in encounters]
    activity_ids = [a.get("id") for a in activities]

    # Build a lookup:  (activity_id, encounter_id) -> condition
    soa_map: dict[tuple[str, str], str] = {}
    for tl in timelines:
        for inst in tl.get("instances", []):
            enc_id = inst.get("encounterId")
            condition = inst.get("defaultConditionId", "mandatory")
            for act_id in inst.get("activityIds", []):
                soa_map[(act_id, enc_id)] = condition

    # Create the table: 1 header row + len(activities) data rows
    num_cols = 1 + len(encounters)  # activity name + one col per encounter
    table = doc.add_table(rows=1 + len(activities), cols=num_cols)
    _set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Activity"
    _set_cell_shading(header_cells[0], "002060")
    for para in header_cells[0].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    for col_idx, enc in enumerate(encounters, start=1):
        cell = header_cells[col_idx]
        cell.text = enc.get("label", enc.get("name", ""))
        _set_cell_shading(cell, "002060")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(8)

    # Data rows
    for row_idx, act in enumerate(activities, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = act.get("name", "")
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)

        for col_idx, enc_id in enumerate(encounter_ids, start=1):
            key = (act.get("id"), enc_id)
            condition = soa_map.get(key)
            if condition:
                mark = "X" if condition == "mandatory" else "C"
                cells[col_idx].text = mark
                for para in cells[col_idx].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(8)

    # Auto-set narrow column widths for encounter columns
    for col_idx in range(1, num_cols):
        for row in table.rows:
            row.cells[col_idx].width = Cm(1.4)

    doc.add_paragraph()
    doc.add_paragraph("X = Mandatory    C = Conditional",
                      style="Normal")
    doc.add_page_break()


def _add_section_1(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Section 1: Introduction."""
    narrative: dict = protocol_db.narrative_sections or {}

    doc.add_heading("1. INTRODUCTION", level=1)

    sec_1_1 = narrative.get("section_1_1", "")
    if sec_1_1:
        doc.add_heading("1.1 Background and Rationale", level=2)
        doc.add_paragraph(sec_1_1)

    sec_1_2 = narrative.get("section_1_2", "")
    if sec_1_2:
        doc.add_heading("1.2 Benefit-Risk Assessment", level=2)
        doc.add_paragraph(sec_1_2)


def _add_section_2(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Section 2: Objectives and Endpoints."""
    doc.add_heading("2. OBJECTIVES AND ENDPOINTS", level=1)

    study_design_data: dict = protocol_db.study_design_data or {}
    objectives: list[dict] = study_design_data.get("objectives", [])

    if not objectives:
        doc.add_paragraph("No objectives have been defined.")
        return

    # Build a table: Objective Level | Objective | Endpoint(s)
    table = doc.add_table(rows=1, cols=3)
    _set_table_borders(table)

    header = table.rows[0].cells
    for idx, label in enumerate(["Level", "Objective", "Endpoint(s)"]):
        header[idx].text = label
        _set_cell_shading(header[idx], "D9E2F3")
        for para in header[idx].paragraphs:
            for run in para.runs:
                run.bold = True

    for obj in objectives:
        lvl = obj.get("level", "")
        level = lvl.get("decode", "") if isinstance(lvl, dict) else str(lvl)
        endpoints = obj.get("endpoints", [])
        ep_text = "\n".join(
            f"- {ep.get('name', '')}: {ep.get('description', '')}"
            for ep in endpoints
        ) if endpoints else "None defined"

        row = table.add_row().cells
        row[0].text = level
        row[1].text = f"{obj.get('name', '')}\n{obj.get('description', '')}"
        row[2].text = ep_text

    doc.add_paragraph()


def _add_section_3(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Section 3: Trial Design."""
    doc.add_heading("3. TRIAL DESIGN", level=1)

    study_design_data: dict = protocol_db.study_design_data or {}

    # 3.1 Description of study design
    doc.add_heading("3.1 Description of Study Design", level=2)

    im = study_design_data.get("interventionModel")
    if isinstance(im, dict):
        doc.add_paragraph(f"Intervention Model: {im.get('decode', 'Not specified')}")

    bs = study_design_data.get("blindingSchema")
    if isinstance(bs, dict):
        doc.add_paragraph(f"Blinding: {bs.get('decode', 'Not specified')}")

    # 3.2 Study arms
    arms: list[dict] = study_design_data.get("studyArms", [])
    if arms:
        doc.add_heading("3.2 Study Arms", level=2)
        table = doc.add_table(rows=1, cols=4)
        _set_table_borders(table)
        for idx, label in enumerate(["Arm Name", "Type", "Description", "Order"]):
            table.rows[0].cells[idx].text = label
            _set_cell_shading(table.rows[0].cells[idx], "D9E2F3")
            for para in table.rows[0].cells[idx].paragraphs:
                for run in para.runs:
                    run.bold = True

        for arm in sorted(arms, key=lambda a: a.get("order", 0)):
            row = table.add_row().cells
            row[0].text = arm.get("name", "")
            arm_type = arm.get("type", {})
            row[1].text = arm_type.get("decode", "") if isinstance(arm_type, dict) else ""
            row[2].text = arm.get("description", "")
            row[3].text = str(arm.get("order", ""))

    # 3.3 Study epochs
    epochs: list[dict] = study_design_data.get("studyEpochs", [])
    if epochs:
        doc.add_heading("3.3 Study Epochs", level=2)
        table = doc.add_table(rows=1, cols=3)
        _set_table_borders(table)
        for idx, label in enumerate(["Epoch Name", "Type", "Order"]):
            table.rows[0].cells[idx].text = label
            _set_cell_shading(table.rows[0].cells[idx], "D9E2F3")
            for para in table.rows[0].cells[idx].paragraphs:
                for run in para.runs:
                    run.bold = True

        for epoch in sorted(epochs, key=lambda e: e.get("order", 0)):
            row = table.add_row().cells
            row[0].text = epoch.get("name", "")
            epoch_type = epoch.get("type", {})
            row[1].text = epoch_type.get("decode", "") if isinstance(epoch_type, dict) else ""
            row[2].text = str(epoch.get("order", ""))

    doc.add_paragraph()


def _add_section_4(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Section 4: Trial Population (eligibility criteria)."""
    doc.add_heading("4. TRIAL POPULATION", level=1)

    study_design_data: dict = protocol_db.study_design_data or {}
    raw_criteria = study_design_data.get("eligibilityCriteria", [])

    # Handle both formats: list with category field, or dict with inclusion/exclusion keys
    if isinstance(raw_criteria, dict):
        inclusion = raw_criteria.get("inclusion", [])
        exclusion = raw_criteria.get("exclusion", [])
    elif isinstance(raw_criteria, list):
        inclusion = sorted(
            [c for c in raw_criteria if c.get("category") == "inclusion"],
            key=lambda c: c.get("order", 0),
        )
        exclusion = sorted(
            [c for c in raw_criteria if c.get("category") == "exclusion"],
            key=lambda c: c.get("order", 0),
        )
    else:
        inclusion = []
        exclusion = []

    if not inclusion and not exclusion:
        doc.add_paragraph("No eligibility criteria have been defined.")
        return

    if inclusion:
        doc.add_heading("4.1 Inclusion Criteria", level=2)
        for crit in inclusion:
            label = crit.get("name", "")
            text = crit.get("text", "")
            doc.add_paragraph(f"{label}: {text}", style="List Bullet")

    if exclusion:
        doc.add_heading("4.2 Exclusion Criteria", level=2)
        for crit in exclusion:
            label = crit.get("name", "")
            text = crit.get("text", "")
            doc.add_paragraph(f"{label}: {text}", style="List Bullet")


def _add_section_5(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Section 5: Trial Interventions."""
    doc.add_heading("5. TRIAL INTERVENTIONS", level=1)

    narrative: dict = protocol_db.narrative_sections or {}
    sec_5 = narrative.get("section_5", "")
    if sec_5:
        doc.add_paragraph(sec_5)

    # Also list interventions from study_design_data if available
    study_design_data: dict = protocol_db.study_design_data or {}
    interventions: list[dict] = study_design_data.get("studyInterventions", [])
    if interventions:
        doc.add_heading("5.1 Investigational Product(s)", level=2)
        for intv in interventions:
            doc.add_paragraph(
                f"{intv.get('name', '')}: {intv.get('description', '')}",
                style="List Bullet",
            )


def _add_section_6(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Section 6: Endpoints (detailed)."""
    doc.add_heading("6. ENDPOINTS", level=1)

    study_design_data: dict = protocol_db.study_design_data or {}
    objectives: list[dict] = study_design_data.get("objectives", [])

    if not objectives:
        doc.add_paragraph("No endpoints have been defined.")
        return

    # Group by objective level
    level_groups: dict[str, list[dict]] = {}
    for obj in objectives:
        lvl = obj.get("level", "Other")
        level_decode = lvl.get("decode", "Other") if isinstance(lvl, dict) else str(lvl) or "Other"
        level_groups.setdefault(level_decode, []).append(obj)

    section_num = 1
    for level_name, objs in level_groups.items():
        doc.add_heading(f"6.{section_num} {level_name} Endpoints", level=2)
        for obj in objs:
            doc.add_paragraph(f"Objective: {obj.get('name', '')}", style="Normal")
            for ep in obj.get("endpoints", []):
                doc.add_paragraph(
                    f"{ep.get('name', '')}: {ep.get('description', '')}",
                    style="List Bullet",
                )
        section_num += 1


def _add_remaining_sections(doc: Document, protocol_db: "ProtocolDB") -> None:
    """Sections 7-11 from narrative_sections."""
    narrative: dict = protocol_db.narrative_sections or {}

    section_map = [
        ("7", "DISCONTINUATION OF TRIAL TREATMENT AND PARTICIPANT WITHDRAWAL",
         "section_7"),
        ("8", "TRIAL ASSESSMENTS AND PROCEDURES", "section_8"),
        ("9", "STATISTICAL CONSIDERATIONS", "section_9"),
        ("10", "SAFETY REPORTING", "section_10"),
        ("11", "OVERSIGHT AND MONITORING", "section_11"),
    ]

    for sec_num, title, key in section_map:
        doc.add_heading(f"{sec_num}. {title}", level=1)
        content = narrative.get(key, "")
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_paragraph(f"[Content for Section {sec_num} to be provided.]")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_protocol_document(protocol_db: "ProtocolDB") -> Document:
    """Generate an ICH M11 Word document from a ProtocolDB row.

    Parameters
    ----------
    protocol_db:
        A SQLAlchemy ``ProtocolDB`` instance loaded from the database.

    Returns
    -------
    docx.Document
        A ``python-docx`` Document object ready to be saved or streamed
        to the caller.
    """
    doc = Document()

    # ---- Global style tweaks -----------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    for heading_level in range(1, 4):
        heading_style_name = f"Heading {heading_level}"
        if heading_style_name in doc.styles:
            h_style = doc.styles[heading_style_name]
            h_style.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    # ---- Build document sections -------------------------------------------
    _add_title_page(doc, protocol_db)
    _add_synopsis(doc, protocol_db)
    _add_schedule_of_activities(doc, protocol_db)
    _add_section_1(doc, protocol_db)
    _add_section_2(doc, protocol_db)
    _add_section_3(doc, protocol_db)
    _add_section_4(doc, protocol_db)
    _add_section_5(doc, protocol_db)
    _add_section_6(doc, protocol_db)
    _add_remaining_sections(doc, protocol_db)

    return doc
